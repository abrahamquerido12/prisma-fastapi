from typing import Dict
import os

import pdfkit
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi_mail import FastMail, MessageSchema, ConnectionConfig, MessageType
from fastapi.responses import JSONResponse
from fastapi import File


conf = ConnectionConfig(
    MAIL_USERNAME = os.environ.get("EMAIL_ADDRESS"),
    MAIL_PASSWORD =  os.environ.get("EMAIL_PASS"),
    MAIL_PORT = 587,
    MAIL_SERVER = "smtp.zoho.com",
    MAIL_SSL_TLS = False,
    MAIL_STARTTLS= True,
    MAIL_FROM = os.environ.get("EMAIL_ADDRESS"),
  
)

import os

def generar_oficio(usuario):
    # Leer el archivo de plantilla
    with open("src/utils/formato.txt", encoding='utf-8') as f:
        plantilla = f.read()
        
    # Reemplazar los campos de la plantilla con los datos del usuario
    datos_oficio = plantilla.format(
        nombre=usuario.profile.name,
        apellido1=usuario.profile.firstLastName,
        apellido2=usuario.profile.secondLastName,
        cargo=usuario.profile.position,
        empresa=usuario.profile.company,
        calle=usuario.profile.street,
        numeroExt=usuario.profile.extNumber,
        numeroInt=usuario.profile.intNumber,
        colonia=usuario.profile.suburb,
        municipio=usuario.profile.city,
        estado=usuario.profile.state,
        codigoPostal=usuario.profile.zipCode,
        telefono=usuario.profile.phone,
        correoElectronico=usuario.email,
        fechaNacimiento=usuario.profile.birthDay,
        edad=usuario.profile.age
    )

    # crear un txt con esta información en src/utils/oficios y regresaremos la ruta del archivo
    ruta_archivo = f"src/utils/oficios/{usuario.profile.name}-{usuario.profile.firstLastName}.txt"
    with open(ruta_archivo, "w") as f:
        f.write(datos_oficio)
        
    
    return ruta_archivo

def generar_oficio_pdf(usuario):
    ruta_archivo = f"src/utils/oficios/{usuario.profile.name}-{usuario.profile.firstLastName}.txt"
    
    #if no file create one
    if not os.path.isfile(ruta_archivo):
        generar_oficio(usuario)

    # Obtener la ruta donde se guardará el archivo PDF
    ruta_pdf = f"src/utils/oficios/{usuario.profile.name}-{usuario.profile.firstLastName}.pdf"


    # Convertir el archivo de texto a HTML
    with open(ruta_archivo, encoding='latin-1') as f:
        text = f.read()
        html = f"""
            <html>
                <head>
                    <style>
                        h1 {{
                            font-size: 24px;
                            text-align: center;
                            margin-top: 20px;
                            margin-bottom: 20px;
                        }}
                        h2 {{
                            font-size: 20px;
                            margin-bottom: 10px;
                        }}
                        p {{
                            font-size: 16px;
                            margin: 0;
                            line-height: 1.5;
                        }}
                        .section {{
                            margin-top: 20px;
                            margin-bottom: 20px;
                        }}
                    </style>
                </head>
                <body>
                    <h1>Oficio de presentación</h1>
                    <div class="section">
                        <h2>Estimado(a) {usuario.profile.name} {usuario.profile.firstLastName}:</h2>
                        <p>Por medio de la presente, nos es grato hacerle llegar la siguiente información sobre su registro con nosotros:</p>
                        <p>Nombre completo: {usuario.profile.name} {usuario.profile.firstLastName} {usuario.profile.secondLastName}</p>
                        <p>Cargo: {usuario.profile.position}</p>
                        <p>Empresa: {usuario.profile.company}</p>
                        <p>Dirección: {usuario.profile.street} {usuario.profile.extNumber} {usuario.profile.intNumber}, Col. {usuario.profile.suburb}, {usuario.profile.city}, {usuario.profile.state}, C.P. {usuario.profile.zipCode}</p>
                        <p>Teléfono: {usuario.profile.phone}</p>
                        <p>Correo electrónico: {usuario.email}</p>
                        <p>Fecha de nacimiento: {usuario.profile.birthDay}</p>
                        <p>Edad: {usuario.profile.age}</p>
                        <p>Agradecemos su interés en nuestros servicios y quedamos a sus órdenes para cualquier consulta o aclaración.</p>
                        <p>Atentamente,</p>
                        <p>El equipo de nuestra empresa</p>
                    </div>
                </body>
            </html>
        """  

    config = pdfkit.configuration(wkhtmltopdf="C:/Program Files/wkhtmltopdf/bin/wkhtmltopdf.exe")
    options = {
    'encoding': 'UTF-8'
    }
    #  Generar el archivo PDF
    pdfkit.from_string(html, ruta_pdf, configuration=config, options=options)

    # Verificar si el archivo PDF fue creado exitosamente
    if os.path.isfile(ruta_pdf):
        return ruta_pdf
    else:
        return {"error": "No se pudo generar el archivo PDF."}
    

def generar_doc (usuario):
    doc = Document()

    title = doc.add_heading('Oficio de presentación', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    saludo = doc.add_paragraph('Estimado(a) ' + usuario.profile.name + ' ' + usuario.profile.firstLastName + ':')
    saludo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    informacion_usuario = doc.add_paragraph()
    informacion_usuario.add_run('Por medio de la presente, nos es grato hacerle llegar la siguiente información sobre su registro con nosotros:\n').bold = True
    informacion_usuario.add_run(f'Nombre completo: {usuario.profile.name} {usuario.profile.firstLastName} {usuario.profile.secondLastName}\n')
    informacion_usuario.add_run(f'Cargo: {usuario.profile.position}\n')
    informacion_usuario.add_run(f'Empresa: {usuario.profile.company}\n')
    informacion_usuario.add_run(f'Dirección: {usuario.profile.street} {usuario.profile.extNumber} {usuario.profile.intNumber}, Col. {usuario.profile.suburb}, {usuario.profile.city}, {usuario.profile.state}, C.P. {usuario.profile.zipCode}\n')
    informacion_usuario.add_run(f'Teléfono: {usuario.profile.phone}\n')
    informacion_usuario.add_run(f'Correo electrónico: {usuario.email}\n')
    informacion_usuario.add_run(f'Fecha de nacimiento: {usuario.profile.birthDay}\n')
    informacion_usuario.add_run(f'Edad: {usuario.profile.age}\n')
    informacion_usuario.alignment = WD_ALIGN_PARAGRAPH.LEFT
    informacion_usuario.paragraph_format.space_before = Pt(12)

    agradecimiento = doc.add_paragraph('Agradecemos su interés en nuestros servicios y quedamos a sus órdenes para cualquier consulta o aclaración.\n')
    agradecimiento.alignment = WD_ALIGN_PARAGRAPH.LEFT
    agradecimiento.paragraph_format.space_before = Pt(12)

    despedida = doc.add_paragraph('Atentamente,\nEl equipo de nuestra empresa\n')
    despedida.alignment = WD_ALIGN_PARAGRAPH.LEFT
    despedida.paragraph_format.space_before = Pt(24)

    ruta_word = f"src/utils/oficios/{usuario.profile.name}-{usuario.profile.firstLastName}.docx"

    doc.save(ruta_word)

    if os.path.isfile(ruta_word):
        return ruta_word
    else:
        return {"error": "No se pudo generar el archivo Word."}
    
async def send_mail(user, file_type: str = "txt"):
    fastmail = FastMail(conf)
    ruta_archivo = generar_oficio(user)

    if file_type == "pdf":
        ruta_archivo = ruta_archivo.replace(".txt", ".pdf")
        generar_oficio_pdf(user)
    elif file_type == "doc":
        ruta_archivo = ruta_archivo.replace(".txt", ".docx")
        generar_doc(user)


    message = MessageSchema(
        subject="Oficio",
        recipients=[user.email],
        body=f"Estimado(a) {user.profile.name} {user.profile.firstLastName},\n\nAdjuntamos el archivo del oficio correspondiente a su registro con nosotros.\n\nAtentamente,\nEl equipo de nuestra empresa",
        subtype=MessageType.plain,
        attachments=[ruta_archivo]
        
    )
    


    try:
        await fastmail.send_message(message)
        return JSONResponse(content={"message": f"El oficio ha sido enviado por correo electrónico a {user.email}."})
    except Exception as e:
        return JSONResponse(content={"message": f"Error al enviar el correo electrónico: {str(e)}"}, status_code=500)